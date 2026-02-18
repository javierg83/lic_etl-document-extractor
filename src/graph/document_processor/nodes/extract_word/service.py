import docx

class ExtractWordService:
    @staticmethod
    def extract_text(file_path: str) -> list:
        """
        Extrae contenido de un Word usando python-docx.
        Retorna una lista de diccionarios con la estructura:
        [{"page": 1, "text": "..."}]
        (Word no tiene paginación real, se retorna todo en 'page': 1)
        """
        print(f"📝 Extrayendo Word con python-docx: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extraer tablas (opcional: se añaden al final o intercaladas si se pudiera determinar posición)
            # Por simplicidad en python-docx, las tablas se acceden por separado.
            # Una estrategia común es iterar el cuerpo del documento, pero docx.Document
            # separa paragraphs y tables. Para mantener orden visual aproximado,
            # iteramos paragraphs. Si se requiere tablas, se pueden añadir al final
            # o intentar una lógica más compleja de iteración de elementos del body.
            # Aquí añadiremos las tablas al final del texto por ahora.
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))

            text_content = "\n".join(full_text)
            
            return [
                {
                    "page": 1,
                    "text": text_content
                }
            ]
            
        except Exception as e:
            print(f"❌ Error leyendo Word {file_path}: {e}")
            raise e
